#!/usr/bin/env python3
"""生成监狱出监管理系统 API 接口文档 (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_header_row(table, texts, bold=True):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2F5496')
        run.font.color.rgb = RGBColor(255, 255, 255)


def add_data_row(table, texts, bold_first=False):
    row = table.add_row()
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if bold_first and i == 0:
            run.bold = True


def set_table_width(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)


def create_field_table(doc, fields, col_widths=None):
    """创建字段表格，fields: [(字段名, 类型, 说明, 必填, 默认值), ...]"""
    headers = ['字段名', '类型', '说明', '必填', '默认值']
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    add_header_row(table, headers)
    for f in fields:
        add_data_row(table, f)
    set_table_width(table)
    return table


def create_api_table(doc, fields):
    """创建API参数表格"""
    headers = ['参数名', '类型', '位置', '必填', '说明']
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    add_header_row(table, headers)
    for f in fields:
        add_data_row(table, f)
    set_table_width(table)
    return table


def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)
    return p


def add_api_section(doc, method, path, title, desc, auth, params, response_example):
    """添加一个API接口章节"""
    add_title(doc, f'{method} {path}', level=3)
    add_para(doc, desc)

    add_para(doc, f'认证方式: {auth}', bold=True)

    if params:
        add_para(doc, '请求参数:', bold=True)
        create_api_table(doc, params)

    add_para(doc, '响应示例:', bold=True)
    add_code(doc, response_example)
    doc.add_paragraph('')


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ========== 封面 ==========
    for _ in range(4):
        doc.add_paragraph('')

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('监狱出监管理系统')
    run.font.size = Pt(28)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = True
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('API 接口文档')
    run.font.size = Pt(22)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    doc.add_paragraph('')

    version_p = doc.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_p.add_run(f'版本: V1.0\n日期: {datetime.date.today().strftime("%Y-%m-%d")}')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ========== 目录 ==========
    add_title(doc, '目录', level=1)
    toc_items = [
        '一、概述',
        '二、通用说明',
        '  2.1 基础URL',
        '  2.2 认证方式',
        '  2.3 通用响应格式',
        '三、数据库表结构',
        '  3.1 用户表 (user_login)',
        '  3.2 出入监记录表 (exit_entry_record)',
        '  3.3 每日统计表 (daily_statistics)',
        '  3.4 历史统计表 (history_statistics)',
        '  3.5 出监类型表 (exit_type)',
        '  3.6 罪犯档案表 (prisoner_archive)',
        '四、接口列表',
        '  4.1 登录认证',
        '  4.2 账号管理',
        '  4.3 出监记录',
        '  4.4 入监记录',
        '  4.5 回监记录',
        '  4.6 记录查询与导出',
        '  4.7 统计分析',
        '  4.8 出监类型管理',
        '  4.9 监狱消息',
        '  4.10 视频管理',
        '  4.11 档案管理',
        '  4.12 数据同步',
        '  4.13 人脸抓拍',
    ]
    for item in toc_items:
        add_para(doc, item)

    doc.add_page_break()

    # ========== 一、概述 ==========
    add_title(doc, '一、概述', level=1)
    add_para(doc, '本文档为监狱出监管理系统的后端API接口文档。系统基于Django + Django REST Framework开发，提供RESTful风格的API接口，用于管理罪犯的出监、入监、回监流程，以及相关的统计分析、视频管理、档案管理和大华人脸识别设备数据同步等功能。')
    add_para(doc, '技术栈: Django 4.x + DRF + Celery + Redis + PostgreSQL + Django Channels')
    add_para(doc, '前端: React + Ant Design v4')

    doc.add_page_break()

    # ========== 二、通用说明 ==========
    add_title(doc, '二、通用说明', level=1)

    add_title(doc, '2.1 基础URL', level=2)
    add_code(doc, 'http://<server_ip>:8000/api/')

    add_title(doc, '2.2 认证方式', level=2)
    add_para(doc, '除登录接口外，所有接口均需要JWT Token认证。')
    add_para(doc, '在请求头中携带:')
    add_code(doc, 'Authorization: Bearer <token>')
    add_para(doc, 'Token通过登录接口获取，有效期由服务端配置决定。')

    add_title(doc, '2.3 通用响应格式', level=2)
    add_para(doc, '所有接口返回统一的JSON格式:')
    add_code(doc, '{\n    "code": 1,       // 1=成功, 0=失败\n    "msg": "操作成功", // 提示信息\n    "data": {},        // 返回数据\n    "num": 10          // 总条数(分页接口)\n}')
    add_para(doc, 'HTTP状态码统一返回200，业务状态通过code字段区分。')

    doc.add_page_break()

    # ========== 三、数据库表结构 ==========
    add_title(doc, '三、数据库表结构', level=1)

    # ----- 3.1 用户表 -----
    add_title(doc, '3.1 用户表 (user_login)', level=2)
    add_para(doc, '系统用户表，继承Django AbstractUser，扩展了角色和监狱信息字段。')
    create_field_table(doc, [
        ('id', 'AutoField', '主键ID', '自动', '-'),
        ('username', 'CharField(150)', '用户名', '是', '-'),
        ('password', 'CharField(128)', '密码(哈希)', '是', '-'),
        ('first_name', 'CharField(150)', '显示名称', '否', '空'),
        ('email', 'CharField(254)', '邮箱', '否', '空'),
        ('is_staff', 'BooleanField', '是否员工', '否', 'False'),
        ('is_active', 'BooleanField', '是否激活', '否', 'True'),
        ('is_superuser', 'BooleanField', '是否超级管理员', '否', 'False'),
        ('date_joined', 'DateTimeField', '注册时间', '自动', '-'),
        ('last_login', 'DateTimeField', '最后登录时间', '自动', '-'),
        ('role', 'CharField(32)', '角色: admin/user', '否', 'user'),
        ('role_name', 'CharField(64)', '角色名称: 管理员/普通用户', '否', '空'),
        ('prison_id', 'CharField(32)', '所属监狱ID', '否', '空'),
        ('prison_name', 'CharField(128)', '所属监狱名称', '否', '空'),
    ])

    doc.add_paragraph('')

    # ----- 3.2 出入监记录表 -----
    add_title(doc, '3.2 出入监记录表 (exit_entry_record)', level=2)
    add_para(doc, '记录罪犯的出监、入监、回监操作，包含各环节人脸照片和签名信息。')
    create_field_table(doc, [
        ('id', 'AutoField', '主键ID', '自动', '-'),
        ('prisoner_no', 'CharField(32)', '罪犯编号', '是', '-'),
        ('prisoner_name', 'CharField(64)', '罪犯姓名', '是', '-'),
        ('prisoner_photo', 'CharField(255)', '罪犯照片路径', '否', '空'),
        ('prison_area', 'CharField(32)', '监区ID', '否', '空'),
        ('prison_area_name', 'CharField(128)', '监区名称', '否', '空'),
        ('type', 'CharField(16)', '类型: exit=出监, entry=入监', '是', '-'),
        ('reason', 'CharField(32)', '出监原因', '否', '空'),
        ('exit_date', 'DateTimeField', '出监日期', '否', 'null'),
        ('entry_date', 'DateTimeField', '入监日期', '否', 'null'),
        ('police_face', 'CharField(255)', '民警人脸照片路径', '否', '空'),
        ('police_name', 'CharField(64)', '民警姓名', '否', '空'),
        ('swat_face', 'CharField(255)', '特警人脸照片路径', '否', '空'),
        ('swat_name', 'CharField(64)', '特警姓名', '否', '空'),
        ('armed_police_signature', 'TextField', '武警签名(base64)', '否', '空'),
        ('armed_police_face', 'CharField(512)', '武警人脸照片路径', '否', 'null'),
        ('armed_police_name', 'CharField(64)', '武警姓名', '否', '空'),
        ('hospital_type', 'CharField(32)', '医院类型', '否', 'null'),
        ('hospital_name', 'CharField(128)', '医院名称', '否', 'null'),
        ('operator_id', 'IntegerField', '操作人ID', '否', 'null'),
        ('operator_name', 'CharField(64)', '操作人姓名', '否', '空'),
        ('status', 'CharField(16)', '状态: processing/completed', '否', 'processing'),
        ('abnormal_reason', 'CharField(255)', '异常原因', '否', '空'),
        ('attachments', 'JSONField', '附件列表', '否', '[]'),
        ('start_time', 'CharField(20)', '开始时间', '否', '空'),
        ('end_time', 'CharField(20)', '结束时间', '否', '空'),
        ('video_url', 'CharField(512)', '录像存储URL', '否', 'null'),
        ('created_at', 'DateTimeField', '创建时间', '自动', '-'),
        ('updated_at', 'DateTimeField', '更新时间', '自动', '-'),
    ])

    doc.add_paragraph('')

    # ----- 3.3 每日统计表 -----
    add_title(doc, '3.3 每日统计表 (daily_statistics)', level=2)
    add_para(doc, '按监区和日期统计每日出监、入监、在监、出工人数等数据。唯一约束: (prison_area, date)。')
    create_field_table(doc, [
        ('id', 'AutoField', '主键ID', '自动', '-'),
        ('prison_area', 'CharField(32)', '监区ID', '是', '-'),
        ('prison_area_name', 'CharField(128)', '监区名称', '否', '空'),
        ('date', 'DateField', '统计日期', '是', '-'),
        ('exit_count', 'IntegerField', '出监总人数', '否', '0'),
        ('entry_count', 'IntegerField', '入监总人数', '否', '0'),
        ('in_prison_count', 'IntegerField', '实时在监人数', '否', '0'),
        ('work_count', 'IntegerField', '出工人数', '否', '0'),
        ('reason_stats', 'JSONField', '出监原因统计(dict)', '否', '{}'),
        ('created_at', 'DateTimeField', '创建时间', '自动', '-'),
        ('updated_at', 'DateTimeField', '更新时间', '自动', '-'),
    ])

    doc.add_paragraph('')

    # ----- 3.4 历史统计表 -----
    add_title(doc, '3.4 历史统计表 (history_statistics)', level=2)
    add_para(doc, '按监区和日期记录历史统计数据，包含各类出监原因的分项计数。')
    create_field_table(doc, [
        ('id', 'AutoField', '主键ID', '自动', '-'),
        ('prison_area', 'CharField(32)', '监区ID', '是', '-'),
        ('prison_area_name', 'CharField(128)', '监区名称', '否', '空'),
        ('date', 'DateField', '统计日期', '是', '-'),
        ('exit_count', 'IntegerField', '出监总人数', '否', '0'),
        ('exit_reason_1', 'IntegerField', '刑满释放人数', '否', '0'),
        ('exit_reason_2', 'IntegerField', '外出就医人数', '否', '0'),
        ('exit_reason_3', 'IntegerField', '外出教育人数', '否', '0'),
        ('exit_reason_4', 'IntegerField', '离监探亲人数', '否', '0'),
        ('exit_reason_5', 'IntegerField', '押回重审人数', '否', '0'),
        ('entry_count', 'IntegerField', '入监总人数', '否', '0'),
        ('created_at', 'DateTimeField', '创建时间', '自动', '-'),
    ])

    doc.add_paragraph('')

    # ----- 3.5 出监类型表 -----
    add_title(doc, '3.5 出监类型表 (exit_type)', level=2)
    add_para(doc, '出监原因类型表，支持树形层级结构（parent自引用）。')
    create_field_table(doc, [
        ('id', 'AutoField', '主键ID', '自动', '-'),
        ('type_name', 'CharField(128)', '出监原因名称', '是', '-'),
        ('parent', 'ForeignKey(self)', '上级出监原因ID', '否', 'null'),
        ('level', 'PositiveIntegerField', '层级(1=一级, 2=二级...)', '否', '1'),
        ('sort_order', 'IntegerField', '排序号', '否', '0'),
        ('status', 'CharField(16)', '状态: active=启用, disabled=停用', '否', 'active'),
        ('created_at', 'DateTimeField', '创建时间', '自动', '-'),
        ('updated_at', 'DateTimeField', '更新时间', '自动', '-'),
    ])

    doc.add_paragraph('')

    # ----- 3.6 罪犯档案表 -----
    add_title(doc, '3.6 罪犯档案表 (prisoner_archive)', level=2)
    add_para(doc, '从公安内网同步的罪犯基本信息和媒体信息，用于出监确认时的人员比对。')
    create_field_table(doc, [
        ('id', 'AutoField', '主键ID', '自动', '-'),
        ('prisoner_no', 'CharField(32)', '罪犯编号(唯一)', '是', '-'),
        ('prisoner_name', 'CharField(64)', '姓名', '否', '空'),
        ('gender', 'CharField(8)', '性别', '否', '空'),
        ('birth_date', 'CharField(20)', '出生日期', '否', '空'),
        ('age', 'IntegerField', '年龄', '否', 'null'),
        ('id_card', 'CharField(32)', '身份证号', '否', '空'),
        ('nation', 'CharField(32)', '民族', '否', '空'),
        ('education', 'CharField(32)', '文化程度', '否', '空'),
        ('marital_status', 'CharField(16)', '婚姻状况', '否', '空'),
        ('native_place', 'CharField(128)', '籍贯', '否', '空'),
        ('address', 'CharField(256)', '家庭地址', '否', '空'),
        ('crime', 'CharField(128)', '罪名', '否', '空'),
        ('sentence', 'CharField(64)', '原判刑期', '否', '空'),
        ('sentence_start', 'CharField(20)', '刑期起日', '否', '空'),
        ('sentence_end', 'CharField(20)', '刑期止日', '否', '空'),
        ('prison_area', 'CharField(64)', '监区', '否', '空'),
        ('room_no', 'CharField(32)', '监室号', '否', '空'),
        ('bed_no', 'CharField(32)', '床号', '否', '空'),
        ('status', 'CharField(32)', '在押状态', '否', '空'),
        ('is_released', 'BooleanField', '是否已释放', '否', 'False'),
        ('entry_date', 'CharField(20)', '入监日期', '否', '空'),
        ('arrest_org', 'CharField(128)', '逮捕机关', '否', '空'),
        ('judgment_org', 'CharField(128)', '判决机关', '否', '空'),
        ('judgment_no', 'CharField(128)', '判决书号', '否', '空'),
        ('basic_info', 'JSONField', '基础信息完整数据(dict)', '否', '{}'),
        ('media_info', 'JSONField', '媒体信息列表(list)', '否', '[]'),
        ('last_synced_photo_url', 'CharField(512)', '上次同步到大华的照片URL', '否', '空'),
        ('synced_at', 'DateTimeField', '同步时间', '自动', '-'),
        ('created_at', 'DateTimeField', '创建时间', '自动', '-'),
    ])

    doc.add_page_break()

    # ========== 四、接口列表 ==========
    add_title(doc, '四、接口列表', level=1)

    # ----- 4.1 登录认证 -----
    add_title(doc, '4.1 登录认证', level=2)

    add_api_section(doc,
        method='POST',
        path='/api/user_login/',
        title='用户登录',
        desc='用户通过用户名和密码登录系统，返回JWT Token。',
        auth='无需认证 (AllowAny)',
        params=[
            ('username', 'string', 'Body', '是', '用户名'),
            ('password', 'string', 'Body', '是', '密码'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "登录成功",\n    "data": {\n        "id": 1,\n        "username": "admin",\n        "name": "管理员",\n        "role": "admin",\n        "role_name": "管理员",\n        "prison_id": "001",\n        "prison_name": "XX监狱",\n        "token": "eyJhbGciOiJIUzI1NiIs..." \n    }\n}'
    )

    # ----- 4.2 账号管理 -----
    add_title(doc, '4.2 账号管理', level=2)
    add_para(doc, '所有账号管理接口仅管理员(role=admin)可访问。')

    add_api_section(doc,
        method='GET',
        path='/api/account/account_list/',
        title='获取账号列表',
        desc='获取所有系统用户账号列表。',
        auth='JWT + 管理员权限',
        params=[],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": [\n        {\n            "id": 1,\n            "username": "admin",\n            "name": "管理员",\n            "role": "admin",\n            "role_name": "管理员",\n            "prison_id": "001",\n            "prison_name": "XX监狱",\n            "status": "active"\n        }\n    ],\n    "num": 1\n}'
    )

    add_api_section(doc,
        method='POST',
        path='/api/account/account_add/',
        title='新增账号',
        desc='创建新的系统用户账号。密码默认为123456。',
        auth='JWT + 管理员权限',
        params=[
            ('username', 'string', 'Body', '是', '用户名'),
            ('password', 'string', 'Body', '否', '密码，默认123456'),
            ('name', 'string', 'Body', '否', '显示名称'),
            ('role', 'string', 'Body', '否', '角色: admin/user，默认user'),
            ('prison_id', 'string', 'Body', '否', '所属监狱ID'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "创建成功",\n    "data": { ... }\n}'
    )

    add_api_section(doc,
        method='POST',
        path='/api/account/account_update/',
        title='更新账号',
        desc='更新指定账号的信息。admin用户不可修改角色。',
        auth='JWT + 管理员权限',
        params=[
            ('id', 'integer', 'Body', '是', '账号ID'),
            ('name', 'string', 'Body', '否', '显示名称'),
            ('role', 'string', 'Body', '否', '角色'),
            ('prison_id', 'string', 'Body', '否', '所属监狱ID'),
            ('password', 'string', 'Body', '否', '新密码(不传则不修改)'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "更新成功",\n    "data": { ... }\n}'
    )

    add_api_section(doc,
        method='POST',
        path='/api/account/reset_password/',
        title='重置密码',
        desc='重置指定账号的密码，不传密码则默认重置为123456。',
        auth='JWT + 管理员权限',
        params=[
            ('id', 'integer', 'Body', '是', '账号ID'),
            ('password', 'string', 'Body', '否', '新密码，默认123456'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "重置成功",\n    "data": { ... }\n}'
    )

    add_api_section(doc,
        method='POST',
        path='/api/account/account_delete/',
        title='删除账号',
        desc='删除指定账号。admin用户不可删除。',
        auth='JWT + 管理员权限',
        params=[
            ('id', 'integer', 'Body', '是', '账号ID'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "删除成功",\n    "data": null\n}'
    )

    # ----- 4.3 出监记录 -----
    add_title(doc, '4.3 出监记录', level=2)

    add_api_section(doc,
        method='POST',
        path='/api/exit_record/submit/',
        title='提交出监记录',
        desc='提交罪犯出监记录，包含民警、特警、武警人脸照片和签名。支持multipart/form-data格式上传文件。出监原因为"外出就医"时hospital_name必填。提交成功后自动生成出监视频。',
        auth='JWT认证',
        params=[
            ('prisoner_no', 'string', 'Body', '是', '罪犯编号'),
            ('prisoner_name', 'string', 'Body', '是', '罪犯姓名'),
            ('prison_area', 'string', 'Body', '是', '监区ID(自动转名称)'),
            ('exit_date', 'date', 'Body', '是', '出监日期 YYYY-MM-DD'),
            ('reason', 'string', 'Body', '是', '出监原因ID(自动转名称)'),
            ('police_face', 'base64', 'Body', '是', '民警人脸照片(base64)'),
            ('swat_face', 'base64', 'Body', '是', '特警人脸照片(base64)'),
            ('armed_police_signature', 'base64', 'Body', '是', '武警签名(base64)'),
            ('prisoner_photo', 'base64', 'Body', '否', '罪犯照片'),
            ('police_name', 'string', 'Body', '否', '民警姓名'),
            ('swat_name', 'string', 'Body', '否', '特警姓名'),
            ('armed_police_face', 'base64', 'Body', '否', '武警人脸照片(base64)'),
            ('hospital_name', 'string', 'Body', '否', '医院名称(外出就医时必填)'),
            ('start_time', 'datetime', 'Body', '否', '开始时间'),
            ('end_time', 'datetime', 'Body', '否', '结束时间'),
            ('attachments', 'file[]', 'Body', '否', '附件文件列表'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "提交成功",\n    "data": {\n        "id": 1,\n        "prisoner_no": "10001",\n        "prisoner_name": "张三",\n        "type": "exit",\n        "status": "completed",\n        "video_url": null,\n        ...\n    }\n}'
    )

    # ----- 4.4 入监记录 -----
    add_title(doc, '4.4 入监记录', level=2)

    add_api_section(doc,
        method='POST',
        path='/api/entry_record/submit/',
        title='提交入监记录',
        desc='提交罪犯入监记录。系统会自动关联该罪犯最近一次出监记录，并调整相关统计数据。如果出监原因为"刑满释放"，会重置罪犯档案的释放状态。',
        auth='JWT认证',
        params=[
            ('prisoner_no', 'string', 'Body', '是', '罪犯编号'),
            ('prisoner_name', 'string', 'Body', '是', '罪犯姓名'),
            ('prison_area', 'string', 'Body', '是', '监区ID'),
            ('entry_date', 'date', 'Body', '是', '入监日期 YYYY-MM-DD'),
            ('police_face', 'base64', 'Body', '是', '民警人脸照片(base64)'),
            ('prisoner_photo', 'base64', 'Body', '否', '罪犯照片'),
            ('police_name', 'string', 'Body', '否', '民警姓名'),
            ('entry_status', 'string', 'Body', '否', '入监状态，默认normal'),
            ('abnormal_reason', 'string', 'Body', '否', '异常原因'),
            ('start_time', 'datetime', 'Body', '否', '开始时间'),
            ('end_time', 'datetime', 'Body', '否', '结束时间'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "提交成功",\n    "data": {\n        "id": 2,\n        "prisoner_no": "10001",\n        "type": "entry",\n        "status": "completed",\n        ...\n    }\n}'
    )

    # ----- 4.5 回监记录 -----
    add_title(doc, '4.5 回监记录', level=2)

    add_api_section(doc,
        method='POST',
        path='/api/return_record/submit/',
        title='提交回监记录',
        desc='提交罪犯回监记录（如外出就医、外出教育、离监探亲后的返回）。参数与入监记录相同。',
        auth='JWT认证',
        params=[
            ('prisoner_no', 'string', 'Body', '是', '罪犯编号'),
            ('prisoner_name', 'string', 'Body', '是', '罪犯姓名'),
            ('prison_area', 'string', 'Body', '是', '监区ID'),
            ('entry_date', 'date', 'Body', '是', '回监日期 YYYY-MM-DD'),
            ('police_face', 'base64', 'Body', '是', '民警人脸照片(base64)'),
            ('prisoner_photo', 'base64', 'Body', '否', '罪犯照片'),
            ('police_name', 'string', 'Body', '否', '民警姓名'),
            ('entry_status', 'string', 'Body', '否', '回监状态，默认normal'),
            ('abnormal_reason', 'string', 'Body', '否', '异常原因'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "提交成功",\n    "data": { ... }\n}'
    )

    # ----- 4.6 记录查询与导出 -----
    add_title(doc, '4.6 记录查询与导出', level=2)

    add_api_section(doc,
        method='GET',
        path='/api/record/list/',
        title='查询记录列表',
        desc='分页查询出入监记录，支持多种筛选条件。非管理员用户只能查看自己监狱的数据。',
        auth='JWT认证',
        params=[
            ('type', 'string', 'Query', '否', '记录类型: exit/entry'),
            ('start_timestamp', 'string', 'Query', '否', '开始时间戳'),
            ('end_timestamp', 'string', 'Query', '否', '结束时间戳'),
            ('prison_area', 'string', 'Query', '否', '监区ID'),
            ('prisoner_name', 'string', 'Query', '否', '罪犯姓名(模糊)'),
            ('prisoner_no', 'string', 'Query', '否', '罪犯编号(模糊)'),
            ('reason', 'string', 'Query', '否', '出监原因'),
            ('page', 'integer', 'Query', '否', '页码，默认1'),
            ('limit', 'integer', 'Query', '否', '每页条数，默认10'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": [\n        {\n            "id": 1,\n            "prisoner_no": "10001",\n            "prisoner_name": "张三",\n            "type": "exit",\n            "reason": "刑满释放",\n            "police_face": "http://...",\n            "video_url": "http://...",\n            "created_at": "2026-07-09 10:30:00",\n            ...\n        }\n    ],\n    "num": 100\n}'
    )

    add_api_section(doc,
        method='GET',
        path='/api/record/export/',
        title='导出记录',
        desc='导出出入监记录(最多5000条)，筛选参数与列表接口相同。',
        auth='JWT认证',
        params=[
            ('type', 'string', 'Query', '否', '记录类型: exit/entry'),
            ('start_timestamp', 'string', 'Query', '否', '开始时间戳'),
            ('end_timestamp', 'string', 'Query', '否', '结束时间戳'),
            ('prison_area', 'string', 'Query', '否', '监区ID'),
            ('prisoner_name', 'string', 'Query', '否', '罪犯姓名'),
            ('prisoner_no', 'string', 'Query', '否', '罪犯编号'),
            ('reason', 'string', 'Query', '否', '出监原因'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": [ ... ]\n}'
    )

    # ----- 4.7 统计分析 -----
    add_title(doc, '4.7 统计分析', level=2)

    add_api_section(doc,
        method='GET',
        path='/api/statistics/realtime/',
        title='实时统计',
        desc='获取实时统计数据，包括各监区在监人数、今日出监/入监人数、年度出监人数、各类出监原因统计等。数据来源: 罪犯档案表(在监人数) + 每日统计表(今日数据) + 出入监记录表(年度数据)。',
        auth='JWT认证',
        params=[],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": {\n        "total": {\n            "exit_count": 5,\n            "entry_count": 3,\n            "in_prison_count": 2400,\n            "yearly_exit_count": 120,\n            "reasons": [\n                {"name": "刑满释放", "count": 80},\n                {"name": "外出就医", "count": 20},\n                ...\n            ]\n        },\n        "by_area": [\n            {\n                "prison_area": "001",\n                "prison_area_name": "一监区",\n                "in_prison_count": 300,\n                "exit_count": 2,\n                "entry_count": 1,\n                "yearly_exit_count": 30,\n                "reasons": [ ... ]\n            }\n        ]\n    }\n}'
    )

    add_api_section(doc,
        method='GET',
        path='/api/statistics/work/',
        title='出工统计',
        desc='获取今日出工统计数据，按监区统计出工人数、在监人数、出监/入监人数。',
        auth='JWT认证',
        params=[],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": {\n        "total": {\n            "work_count": 1800,\n            "in_prison_count": 2400,\n            "exit_count": 5,\n            "entry_count": 3\n        },\n        "by_area": [\n            {\n                "prison_area": "001",\n                "prison_area_name": "一监区",\n                "work_count": 250,\n                "in_prison_count": 300,\n                "exit_count": 2,\n                "entry_count": 1\n            }\n        ]\n    }\n}'
    )

    # ----- 4.8 出监类型管理 -----
    add_title(doc, '4.8 出监类型管理', level=2)

    add_api_section(doc,
        method='GET',
        path='/api/exit_type/list/',
        title='获取出监类型列表',
        desc='获取所有出监类型，支持按名称模糊搜索。',
        auth='JWT认证',
        params=[
            ('type_name', 'string', 'Query', '否', '类型名称(模糊搜索)'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": [\n        {\n            "id": 1,\n            "key": 1,\n            "type_name": "刑满释放",\n            "parent_id": null,\n            "parent_name": null,\n            "level": 1,\n            "sort_order": 1,\n            "status": "active",\n            "children": []\n        }\n    ],\n    "num": 5\n}'
    )

    add_api_section(doc,
        method='POST',
        path='/api/exit_type/add/',
        title='新增出监类型',
        desc='新增出监原因类型，支持设置上级类型形成树形结构。层级自动计算。',
        auth='JWT认证',
        params=[
            ('type_name', 'string', 'Body', '是', '类型名称'),
            ('parent_id', 'integer', 'Body', '否', '上级类型ID'),
            ('sort_order', 'integer', 'Body', '否', '排序号，默认0'),
            ('status', 'string', 'Body', '否', '状态: active/disabled，默认active'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "创建成功",\n    "data": { ... }\n}'
    )

    add_api_section(doc,
        method='POST',
        path='/api/exit_type/update/',
        title='更新出监类型',
        desc='更新出监类型信息。修改层级时会自动更新所有子类型层级。',
        auth='JWT认证',
        params=[
            ('id', 'integer', 'Body', '是', '类型ID'),
            ('type_name', 'string', 'Body', '否', '类型名称'),
            ('parent_id', 'integer', 'Body', '否', '上级类型ID'),
            ('sort_order', 'integer', 'Body', '否', '排序号'),
            ('status', 'string', 'Body', '否', '状态'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "更新成功",\n    "data": { ... }\n}'
    )

    add_api_section(doc,
        method='POST',
        path='/api/exit_type/delete/',
        title='删除出监类型',
        desc='删除指定出监类型。',
        auth='JWT认证',
        params=[
            ('id', 'integer', 'Body', '是', '类型ID'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "删除成功",\n    "data": null\n}'
    )

    # ----- 4.9 监狱消息 -----
    add_title(doc, '4.9 监狱消息', level=2)

    add_api_section(doc,
        method='GET',
        path='/api/record/prison_messages/',
        title='获取监狱消息',
        desc='获取今日刑满释放的罪犯列表（sentence_end为今日且is_released=False）。非管理员只能查看自己监区的消息。',
        auth='JWT认证',
        params=[
            ('page', 'integer', 'Query', '否', '页码，默认1'),
            ('limit', 'integer', 'Query', '否', '每页条数，默认20'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": [\n        {\n            "prisoner_no": "10001",\n            "prisoner_name": "张三",\n            "prison_area": "一监区",\n            "sentence_end": "2026-07-09",\n            ...\n        }\n    ],\n    "num": 3\n}'
    )

    # ----- 4.10 视频管理 -----
    add_title(doc, '4.10 视频管理', level=2)

    add_api_section(doc,
        method='GET',
        path='/api/video/stream_url/',
        title='获取视频流地址',
        desc='根据时间范围和摄像头获取视频。优先返回缓存视频；如无缓存且传入record_id则异步生成(Celery)；否则同步生成。',
        auth='JWT认证',
        params=[
            ('start_time', 'string', 'Query', '是', '开始时间(ISO格式)'),
            ('end_time', 'string', 'Query', '是', '结束时间(ISO格式)'),
            ('camera', 'integer', 'Query', '否', '摄像头索引，默认0'),
            ('record_id', 'integer', 'Query', '否', '记录ID(传入则异步生成)'),
        ],
        response_example='// 缓存命中:\n{\n    "code": 1,\n    "msg": "success",\n    "data": {\n        "url": "/media/videos/xxx.mp4?v=123456",\n        "filename": "xxx.mp4",\n        "camera_name": "大门摄像头",\n        "channel": "1",\n        "is_live": false,\n        "is_cached": true\n    }\n}\n\n// 异步生成:\n{\n    "code": 1,\n    "msg": "视频生成中",\n    "data": {\n        "task_id": "abc-123",\n        "status": "pending",\n        "camera_name": "大门摄像头",\n        "channel": "1"\n    }\n}'
    )

    add_api_section(doc,
        method='GET',
        path='/api/video/task_status/',
        title='查询视频生成任务状态',
        desc='查询异步视频生成任务的状态。',
        auth='JWT认证',
        params=[
            ('task_id', 'string', 'Query', '是', '任务ID'),
        ],
        response_example='// 生成中:\n{\n    "code": 1,\n    "msg": "生成中",\n    "data": {"status": "progress", "task_id": "abc-123"}\n}\n\n// 完成:\n{\n    "code": 1,\n    "msg": "完成",\n    "data": {"status": "success", "task_id": "abc-123", "url": "/media/videos/xxx.mp4"}\n}'
    )

    add_api_section(doc,
        method='GET',
        path='/api/video/camera_list/',
        title='获取摄像头列表',
        desc='获取所有已配置的摄像头信息。',
        auth='JWT认证',
        params=[],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": [\n        {"index": 0, "name": "大门摄像头", "channel": "1", "enabled": true},\n        {"index": 1, "name": "监区摄像头", "channel": "2", "enabled": true}\n    ]\n}'
    )

    # ----- 4.11 档案管理 -----
    add_title(doc, '4.11 档案管理', level=2)

    add_api_section(doc,
        method='GET',
        path='/api/archive/list/',
        title='获取档案列表',
        desc='分页查询罪犯档案，支持按编号、姓名、监区、罪名筛选。返回数据包含基础信息和媒体信息(照片URL已自动修正)。',
        auth='JWT认证',
        params=[
            ('prisoner_no', 'string', 'Query', '否', '罪犯编号(模糊)'),
            ('prisoner_name', 'string', 'Query', '否', '姓名(模糊)'),
            ('prison_area', 'string', 'Query', '否', '监区(模糊)'),
            ('crime', 'string', 'Query', '否', '罪名(模糊)'),
            ('page', 'integer', 'Query', '否', '页码，默认1'),
            ('page_size', 'integer', 'Query', '否', '每页条数，默认10'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": [\n        {\n            "prisoner_no": "10001",\n            "prisoner_name": "张三",\n            "gender": "男",\n            "age": 35,\n            "crime": "盗窃罪",\n            "prison_area": "一监区",\n            "mtxx": [...],\n            "synced_at": "2026-07-09 08:00:00",\n            ...\n        }\n    ],\n    "num": 2400\n}'
    )

    add_api_section(doc,
        method='GET',
        path='/api/archive/detail/',
        title='获取档案详情',
        desc='根据罪犯编号获取档案详细信息。',
        auth='JWT认证',
        params=[
            ('prisoner_no', 'string', 'Query', '是', '罪犯编号'),
        ],
        response_example='{\n    "code": 1,\n    "msg": "success",\n    "data": {\n        "prisoner_no": "10001",\n        "prisoner_name": "张三",\n        ...\n    }\n}'
    )

    # ----- 4.12 数据同步 -----
    add_title(doc, '4.12 数据同步', level=2)

    add_api_section(doc,
        method='POST',
        path='/api/sync/start/',
        title='启动数据同步',
        desc='手动启动从公安内网同步罪犯数据到本地数据库和大华门禁设备的任务(Celery异步)。',
        auth='JWT认证',
        params=[],
        response_example='{\n    "code": 1,\n    "msg": "同步任务已启动",\n    "data": {\n        "task_id": "abc-123"\n    }\n}'
    )

    add_api_section(doc,
        method='GET',
        path='/api/sync/status/',
        title='查询同步状态',
        desc='查询数据同步任务的进度状态。',
        auth='JWT认证',
        params=[
            ('task_id', 'string', 'Query', '是', '任务ID'),
        ],
        response_example='// 进行中:\n{\n    "code": 1,\n    "msg": "success",\n    "data": {\n        "state": "PROGRESS",\n        "current": 50,\n        "total": 100,\n        "step": "syncing",\n        "message": "正在同步 50/100...",\n        "percent": 50\n    }\n}\n\n// 完成:\n{\n    "code": 1,\n    "msg": "success",\n    "data": {\n        "state": "SUCCESS",\n        "current": 100,\n        "total": 100,\n        "step": "done",\n        "message": "同步完成",\n        "percent": 100\n    }\n}'
    )

    # ----- 4.13 人脸抓拍 -----
    add_title(doc, '4.13 人脸抓拍', level=2)

    add_api_section(doc,
        method='GET',
        path='/api/snapshot/',
        title='实时人脸抓拍',
        desc='通过大华智能事件设备(10.2.48.223)的snapshot.cgi接口进行实时抓拍，返回base64编码的JPEG图片。用于出监确认流程中武警人脸拍照。',
        auth='JWT认证',
        params=[
            ('channel', 'string', 'Query', '否', '摄像头通道号，默认"1"'),
        ],
        response_example='{\n    "code": 1,\n    "data": {\n        "image_base64": "/9j/4AAQSkZJRgABAQAAAQ..." \n    }\n}'
    )

    doc.add_page_break()

    # ========== 附录 ==========
    add_title(doc, '附录', level=1)

    add_title(doc, 'A. 出监原因类型', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_header_row(table, ['序号', '原因编码', '说明'])
    reasons = [
        ('1', '刑满释放', '刑期届满释放，会标记罪犯档案为已释放'),
        ('2', '外出就医', '需要外出到医院就诊，hospital_name必填'),
        ('3', '外出教育', '外出参加教育活动'),
        ('4', '离监探亲', '离监探亲'),
        ('5', '押回重审', '押回重新审理'),
    ]
    for r in reasons:
        add_data_row(table, r)
    set_table_width(table)

    doc.add_paragraph('')

    add_title(doc, 'B. 大华设备API', level=2)
    add_para(doc, '系统集成了大华(Dahua)门禁和人脸识别设备，用于:')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_header_row(table, ['设备IP', '用途', 'API'])
    devices = [
        ('10.2.48.224', '罪犯人脸门禁', 'AccessUser.cgi / AccessFace.cgi'),
        ('10.2.48.223', '智能事件(民警/特警/武警拍照)', 'snapshot.cgi / eventManager.cgi'),
    ]
    for d in devices:
        add_data_row(table, d)
    set_table_width(table)

    doc.add_paragraph('')
    add_para(doc, '大华人脸同步使用增量同步机制，通过比对last_synced_photo_url字段判断照片是否需要更新。')

    add_title(doc, 'C. 管理命令', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_header_row(table, ['命令', '说明'])
    cmds = [
        ('python manage.py sync_prisoner_data --real-api --dahua', '同步罪犯数据 + 同步人脸到大华'),
        ('python manage.py sync_dahua_faces', '仅同步人脸到大华(增量)'),
        ('python manage.py sync_dahua_faces --full', '全量同步: 清空设备 → 插入用户 → 插入人脸'),
    ]
    for c in cmds:
        add_data_row(table, c)
    set_table_width(table)

    # 保存
    output_path = '/Users/ran/Documents/work/prison_facial_system/server/API接口文档.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    main()
